#!/usr/bin/env python
# encoding: utf-8
# Author: guoxudong
import json
import sys

from docx import Document
from docx.shared import Inches, Pt

reload(sys)
sys.setdefaultencoding('utf-8')


# 不改变样式替换模板中的内容
def createWord(data):
    # 数据转化为json格式
    info = json.loads(data)
    print info

    def replace_text(table_info,old_text, new_text):
        for p in table_info.paragraphs:
            if old_text in p.text:
                inline = p.runs
                for i in inline:
                    if old_text in i.text:
                        text = i.text.replace(old_text, new_text)
                        i.text = text

    # 打开模板文件
    document = Document('resTemplate/word_template2.docx')
    # 添加图片
    if info.get('imagePath') is not u'':
        imgpath = info.get('imagePath').split('CreateResume')[1]
        if imgpath.startswith((u"\\")):
            imgpath = imgpath[1:]
        else:
            imgpath = imgpath[0:]
        # 最后添加附录
        document.add_paragraph().add_run('照片：'.decode('utf8')).font.size = Pt(15)
        document.add_picture(imgpath, width=Inches(1.5), height=Inches(1.5))
    # 替换内容
    # replace_text('name', '姓名'.decode('utf8'))  # 如果替换的字符为中文，则需要进行解码
    # 操作表格
    tables = document.tables
    #基本信息
    table_info_1 = tables[0].table.cell(1,1)
    table_info_2 = tables[0].table.cell(1,2)
    name = info.get('name')
    replace_text(table_info_1,'name',name)
    replace_text(table_info_1,'sex',info.get('sex'))
    site = info.get('province') + ',' + info.get('city') + ',' + info.get('district')
    replace_text(table_info_1,'site',site)
    replace_text(table_info_1, 'birthday', info.get('birthday'))
    replace_text(table_info_2, 'wechat', info.get('wechat'))
    replace_text(table_info_2, 'phone', info.get('phone'))
    replace_text(table_info_2, 'email', info.get('email'))
    replace_text(table_info_2, 'jobpost', info.get('jobpost'))
    # 经历
    table_info_3 = tables[0].table
    table_info_3.cell(3, 1).text = info.get('post')
    table_info_3.cell(4, 1).text = info.get('duration') + '年'
    price = info.get('price_min') + '-' + info.get('price_max')
    table_info_3.cell(4, 3).text = price
    jobinfos = info.get('jobinfo')
    for i in jobinfos:
        time = i.get('time')
        jobname = i.get('jobname')
        jobtype = i.get('jobtype')
        jobcontent = i.get('jobcontent')
        table_info_3.cell(5, 1).add_paragraph().add_run('时间段：' + time).font.bold = True
        table_info_3.cell(5, 1).add_paragraph().add_run('职业类型：' + jobtype).font.size = Pt(10)
        table_info_3.cell(5, 1).add_paragraph().add_run('职业名称：' + jobname).font.size = Pt(10)
        table_info_3.cell(5, 1).add_paragraph().add_run('工作内容：' + jobcontent).font.size = Pt(10)
        table_info_3.cell(5, 1).add_paragraph('')
    proinfo = info.get('proinfo')
    for i in proinfo:
        time = i.get('time')
        jobname = i.get('jobname')
        jobtype = i.get('jobtype')
        jobcontent = i.get('jobcontent')
        table_info_3.cell(6, 1).add_paragraph().add_run('时间段：' + time).font.bold = True
        table_info_3.cell(6, 1).add_paragraph().add_run('项目角色：' + jobtype).font.size = Pt(10)
        table_info_3.cell(6, 1).add_paragraph().add_run('项目名称：' + jobname).font.size = Pt(10)
        table_info_3.cell(6, 1).add_paragraph().add_run('项目描述：' + jobcontent).font.size = Pt(10)
        table_info_3.cell(6, 1).add_paragraph('')
    # 背景
    table_info_3.cell(8, 1).text = info.get('school')
    table_info_3.cell(9, 1).text = info.get('major')
    table_info_3.cell(9, 3).text = info.get('edu')
    # 外语
    if info.get('cet6') == 'on':
        rank = 'CTE6'
    elif info.get('cet4') == 'on':
        rank = 'CET4'
    else:
        rank = ''
    table_info_3.cell(11, 1).text = rank
    # 自我描述
    introduce = info.get('introduce')
    table_info_3.cell(13, 0).text = introduce

    # 基本信息
    # tables[0].cell(0, 1).text = info.get('name')
    # tables[0].cell(0, 3).text = info.get('sex')
    # site = info.get('province') + ',' + info.get('city') + ',' + info.get('district')
    # tables[0].cell(1, 1).text = site
    # tables[0].cell(1, 3).text = info.get('birthday')
    # tables[0].cell(2, 1).text = info.get('phone')
    # tables[0].cell(2, 3).text = info.get('email')
    # tables[0].cell(3, 1).text = info.get('wechat')
    # tables[0].cell(3, 3).text = info.get('jobpost')
    # # 经历
    # tables[2].cell(0, 1).text = info.get('post')
    # tables[2].cell(1, 1).text = info.get('duration') + '年'
    # price = info.get('price_min') + '-' + info.get('price_max')
    # tables[2].cell(1, 3).text = price
    # jobinfos = info.get('jobinfo')
    # for i in jobinfos:
    #     time = i.get('time')
    #     jobname = i.get('jobname')
    #     jobtype = i.get('jobtype')
    #     jobcontent = i.get('jobcontent')
    #     tables[2].cell(2, 1).add_paragraph().add_run('时间段：' + time).font.bold = True
    #     tables[2].cell(2, 1).add_paragraph().add_run('职业类型：' + jobtype).font.size = Pt(10)
    #     tables[2].cell(2, 1).add_paragraph().add_run('职业名称：' + jobname).font.size = Pt(10)
    #     tables[2].cell(2, 1).add_paragraph().add_run('工作内容：' + jobcontent).font.size = Pt(10)
    #     tables[2].cell(2, 1).add_paragraph('')
    # proinfo = info.get('proinfo')
    # for i in proinfo:
    #     time = i.get('time')
    #     jobname = i.get('jobname')
    #     jobtype = i.get('jobtype')
    #     jobcontent = i.get('jobcontent')
    #     tables[2].cell(3, 1).add_paragraph().add_run('时间段：' + time).font.bold = True
    #     tables[2].cell(3, 1).add_paragraph().add_run('项目角色：' + jobtype).font.size = Pt(10)
    #     tables[2].cell(3, 1).add_paragraph().add_run('项目名称：' + jobname).font.size = Pt(10)
    #     tables[2].cell(3, 1).add_paragraph().add_run('项目描述：' + jobcontent).font.size = Pt(10)
    #     tables[2].cell(3, 1).add_paragraph('')
    # # 背景
    # tables[4].cell(0, 1).text = info.get('school')
    # tables[4].cell(1, 1).text = info.get('major')
    # tables[4].cell(1, 3).text = info.get('edu')
    # # 外语
    # if info.get('cet6') == 'on':
    #     rank = 'CTE6'
    # elif info.get('cet4') == 'on':
    #     rank = 'CET4'
    # else:
    #     rank = ''
    # tables[6].cell(0, 1).text = rank
    # # 自我描述
    # introduce = info.get('introduce')

    # 保存
    # word_name = 'media/word/'+name+'简历.docx'
    document.save('media/word/result.docx')
    return 200

# def wordToHtml():
#     html_parser = HTMLParser.HTMLParser()
#     html = convert('resTemplate/demo.docx')  # 使用docx2html模块将docx文件转成html串，随后你想干嘛都行
#     html = html_parser.unescape(html)
#     print html
#     return html

# if __name__ == '__main__':
#     createWord('1')
